#!/usr/bin/env python3
"""
Markdown to DOCX Converter for LegalAI Training Report
Converts the training report from Markdown to Word DOCX format
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_document_styles(doc):
    """Create custom styles for the document"""
    
    # Title style
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(20)
    title_font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)

    # Heading 1 style
    h1_style = doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    h1_font = h1_style.font
    h1_font.name = 'Arial'
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)

    # Heading 2 style
    h2_style = doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
    h2_font = h2_style.font
    h2_font.name = 'Arial'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_style.paragraph_format.space_before = Pt(10)
    h2_style.paragraph_format.space_after = Pt(4)

    # Heading 3 style
    h3_style = doc.styles.add_style('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
    h3_font = h3_style.font
    h3_font.name = 'Arial'
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_style.paragraph_format.space_before = Pt(8)
    h3_style.paragraph_format.space_after = Pt(4)

    # Normal text style
    normal_style = doc.styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.15

    # Code style
    code_style = doc.styles.add_style('CustomCode', WD_STYLE_TYPE.PARAGRAPH)
    code_font = code_style.font
    code_font.name = 'Courier New'
    code_font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)

def process_markdown_line(line, doc):
    """Process individual markdown line and add to document"""
    line = line.strip()
    
    if not line:
        doc.add_paragraph()
        return
    
    # Main title (single #)
    if line.startswith('# ') and not line.startswith('## '):
        title = line[2:].strip()
        # Remove emoji if present
        title = re.sub(r'📘\s*', '', title)
        p = doc.add_paragraph(title, style='CustomTitle')
        return
    
    # Heading 1 (##)
    elif line.startswith('## ') and not line.startswith('### '):
        heading = line[3:].strip()
        p = doc.add_paragraph(heading, style='CustomHeading1')
        return
    
    # Heading 2 (###)
    elif line.startswith('### ') and not line.startswith('#### '):
        heading = line[4:].strip()
        p = doc.add_paragraph(heading, style='CustomHeading2')
        return
    
    # Heading 3 (####)
    elif line.startswith('#### '):
        heading = line[5:].strip()
        p = doc.add_paragraph(heading, style='CustomHeading3')
        return
    
    # Horizontal rule
    elif line.startswith('---'):
        doc.add_page_break()
        return
    
    # Bold text with checkmarks
    elif line.startswith('- ✅') or line.startswith('- ❌'):
        p = doc.add_paragraph(style='CustomNormal')
        p.add_run(line)
        return
    
    # Lists
    elif line.startswith('- ') or line.startswith('* ') or re.match(r'^\d+\.', line):
        p = doc.add_paragraph(line, style='CustomNormal')
        return
    
    # Code blocks (simplified)
    elif line.startswith('```') or line.startswith('    '):
        if not line.startswith('```'):
            p = doc.add_paragraph(line, style='CustomCode')
        return
    
    # Table detection (simplified)
    elif '|' in line and line.count('|') >= 3:
        # Simple table handling - convert to text for now
        p = doc.add_paragraph(line.replace('|', ' | '), style='CustomNormal')
        return
    
    # Bold text (**text**)
    elif '**' in line:
        p = doc.add_paragraph(style='CustomNormal')
        parts = line.split('**')
        for i, part in enumerate(parts):
            if i % 2 == 0:
                p.add_run(part)
            else:
                p.add_run(part).bold = True
        return
    
    # Regular paragraph
    else:
        p = doc.add_paragraph(line, style='CustomNormal')
        return

def convert_markdown_to_docx(md_file_path, docx_file_path):
    """Convert Markdown file to DOCX"""
    
    # Create new document
    doc = Document()
    
    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Create custom styles
    create_document_styles(doc)
    
    # Read markdown file
    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except Exception as e:
        print(f"Error reading markdown file: {e}")
        return False
    
    # Process each line
    in_code_block = False
    code_content = []
    
    for line in lines:
        line = line.rstrip('\n')
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                if code_content:
                    code_text = '\n'.join(code_content)
                    p = doc.add_paragraph(code_text, style='CustomCode')
                code_content = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            continue
        
        if in_code_block:
            code_content.append(line)
            continue
        
        # Process normal lines
        process_markdown_line(line, doc)
    
    # Save document
    try:
        doc.save(docx_file_path)
        return True
    except Exception as e:
        print(f"Error saving DOCX file: {e}")
        return False

def main():
    """Main function"""
    md_file = 'report.md'
    docx_file = 'LegalAI_Training_Report.docx'
    
    print(f"Converting {md_file} to {docx_file}...")
    
    if convert_markdown_to_docx(md_file, docx_file):
        print(f"✅ Successfully converted to {docx_file}")
    else:
        print("❌ Conversion failed")

if __name__ == "__main__":
    main()
